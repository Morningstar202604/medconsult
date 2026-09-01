"""文档库路由：上传（白名单）/列表/删除；入库即建 RAG 索引。"""
import hashlib
import uuid
from pathlib import Path

from fastapi import APIRouter, File, HTTPException, UploadFile, status
from sqlalchemy import select

from .. import models
from ..config import get_settings
from ..deps import DbDep
from ..deps import CurrentUser
from ..rag import index_doc, remove_doc

router = APIRouter(tags=["library"])

ALLOWED_EXTS = {".txt", ".md", ".pdf", ".docx", ".json", ".csv", ".log", ".xml"}
TEXT_EXTS = {".txt", ".md", ".json", ".csv", ".log", ".xml"}
MAX_BYTES = 32 * 1024 * 1024


def _doc_dir() -> Path:
    s = get_settings()
    base = Path(s.database_url.replace("sqlite:///", "./"))
    d = (base.parent if s.database_url.startswith("sqlite") else Path("./data")) / "documents"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _extract_text(path: Path, ext: str) -> str:
    if ext in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception as e:
            return f"[PDF 解析失败：{e}]"
    if ext == ".docx":
        try:
            import docx
            d = docx.Document(str(path))
            return "\n".join(p.text for p in d.paragraphs if p.text.strip())
        except Exception as e:
            return f"[DOCX 解析失败：{e}]"
    return ""


@router.post("/library/upload")
async def upload(files: list[UploadFile] = File(...), db: DbDep = None, user: CurrentUser = None):
    """批量上传文档；multipart 字段名 files（与前端对齐）。每文件入库并建 RAG 索引。"""
    if not files:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "未收到文件")
    created = []
    for file in files:
        name = file.filename or "未命名.txt"
        ext = Path(name).suffix.lower()
        if ext not in ALLOWED_EXTS:
            raise HTTPException(status.HTTP_400_BAD_REQUEST, f"不支持的文件类型 {ext}")

        data = await file.read()
        if len(data) > MAX_BYTES:
            raise HTTPException(status.HTTP_400_BAD_REQUEST, f"{name} 超过 32MB 上限")
        if not data:
            raise HTTPException(status.HTTP_400_BAD_REQUEST, f"{name} 为空文件")

        storage_name = uuid.uuid4().hex + ext
        dest = _doc_dir() / storage_name
        dest.write_bytes(data)
        text = _extract_text(dest, ext)
        text_path = dest.with_name(dest.stem + ".txt")
        text_path.write_text(text[:60000], encoding="utf-8")
        index_doc(name, text)

        doc = models.Document(name=name, ext=ext, size=len(data),
                              content_hash=hashlib.sha256(data).hexdigest()[:16],
                              storage_name=storage_name, uploaded_by=user.id)
        db.add(doc)
        db.commit()
        db.refresh(doc)
        created.append({"id": doc.id, "name": doc.name, "size": doc.size})
    return {"items": created}


@router.get("/library")
def list_docs(db: DbDep, user: CurrentUser):
    rows = db.scalars(select(models.Document).order_by(models.Document.id.desc())).all()
    return {"items": [{"id": d.id, "name": d.name, "size": d.size,
                       "ext": d.ext, "created_at": d.created_at.isoformat() if d.created_at else None}
                      for d in rows]}


@router.delete("/library/{doc_id}")
def delete_doc(doc_id: int, db: DbDep, user: CurrentUser):
    doc = db.get(models.Document, doc_id)
    if doc is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "文档不存在")
    for p in (_doc_dir() / doc.storage_name, _doc_dir() / (Path(doc.storage_name).stem + ".txt")):
        if p.exists():
            p.unlink()
    remove_doc(doc.name)
    db.delete(doc)
    db.commit()
    return {"ok": True}
